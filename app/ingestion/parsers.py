"""Document parsers for PDF and DOCX extraction."""

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pypdf.

    Preserves structure where possible. Handles corrupt files, password-protected
    PDFs, and extraction errors gracefully.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Clean extracted text with normalized whitespace.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file is corrupt, password-protected, or unreadable.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {file_path}")
    if not path.is_file():
        raise ValueError(f"Path is not a file: {file_path}")

    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required for PDF parsing. Install with: pip install pypdf")

    try:
        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        logger.info("Parsing PDF: %s (%d pages)", path.name, page_count)

        if page_count == 0:
            logger.warning("PDF has no pages: %s", path.name)
            return ""

        text_parts: list[str] = []
        for i, page in enumerate(reader.pages):
            try:
                raw = page.extract_text()
                if raw:
                    text_parts.append(raw)
            except Exception as e:
                logger.warning("Failed to extract page %d from %s: %s", i + 1, path.name, e)

        text = "\n\n".join(text_parts)
        return _clean_text(text)

    except Exception as e:
        err_msg = str(e).lower()
        if "password" in err_msg or "encrypted" in err_msg:
            raise ValueError(f"PDF is password-protected: {file_path}") from e
        raise ValueError(f"Failed to parse PDF {file_path}: {e}") from e


def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx.

    Extracts text from paragraphs and tables, preserving section structure.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Clean extracted text with normalized whitespace.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file is corrupt or unreadable.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    if not path.is_file():
        raise ValueError(f"Path is not a file: {file_path}")

    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

    try:
        doc = Document(str(path))
        section_count = len(doc.sections) if doc.sections else 1
        logger.info("Parsing DOCX: %s (%d sections)", path.name, section_count)

        text_parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        text = "\n\n".join(text_parts)
        return _clean_text(text)

    except Exception as e:
        raise ValueError(f"Failed to parse DOCX {file_path}: {e}") from e


def parse_document(file_path: str) -> str:
    """Parse a document based on its file extension.

    Dispatches to the appropriate parser for .pdf and .docx files.
    Raises ValueError for unsupported formats.

    Args:
        file_path: Path to the document file.

    Returns:
        Clean extracted text.

    Raises:
        ValueError: If the file format is not supported.
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    logger.info("Parsing document: %s (extension: %s)", path.name, ext)

    if ext == ".pdf":
        return parse_pdf(file_path)
    if ext == ".docx":
        return parse_docx(file_path)

    raise ValueError(f"Unsupported document format: {ext}. Supported formats: .pdf, .docx")


def _clean_text(text: str) -> str:
    """Normalize whitespace in extracted text."""
    if not text:
        return ""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
