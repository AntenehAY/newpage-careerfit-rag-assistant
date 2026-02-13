"""Generate sample PDF and DOCX fixtures for ingestion tests."""

from pathlib import Path

FIXTURES_DIR = Path(__file__).parent / "fixtures"
SAMPLE_RESUME_TEXT = (Path(__file__).parent / "fixtures" / "sample_resume.txt").read_text()


def create_sample_pdf() -> Path:
    """Create sample_resume.pdf using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        raise ImportError("reportlab required for PDF fixture. Run: pip install reportlab")

    out_path = FIXTURES_DIR / "sample_resume.pdf"
    c = canvas.Canvas(str(out_path), pagesize=letter)
    width, height = letter
    y = height - 72
    for i, line in enumerate(SAMPLE_RESUME_TEXT.split("\n")):
        if y < 72:
            c.showPage()
            y = height - 72
        c.drawString(72, y, line[:80] if len(line) > 80 else line)
        y -= 14
    c.save()
    return out_path


def create_sample_docx() -> Path:
    """Create sample_jd.docx using python-docx."""
    from docx import Document
    from docx.shared import Pt

    out_path = FIXTURES_DIR / "sample_jd.docx"
    doc = Document()
    doc.add_heading("Senior Backend Engineer", 0)
    doc.add_paragraph(
        "We are seeking a Senior Backend Engineer to join our platform team. "
        "You will design and build scalable APIs and data pipelines."
    )
    doc.add_heading("Requirements", level=1)
    doc.add_paragraph("5+ years of Python experience", style="List Bullet")
    doc.add_paragraph("Experience with AWS or GCP", style="List Bullet")
    doc.add_paragraph("Strong SQL and NoSQL knowledge", style="List Bullet")
    doc.add_heading("Responsibilities", level=1)
    doc.add_paragraph("Design and implement RESTful APIs", style="List Bullet")
    doc.add_paragraph("Collaborate with product and frontend teams", style="List Bullet")
    doc.add_paragraph("Mentor junior engineers", style="List Bullet")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    create_sample_pdf()
    create_sample_docx()
    print("Fixtures created:", list(FIXTURES_DIR.glob("*")))
